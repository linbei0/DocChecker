from io import BytesIO
from pathlib import Path
from time import sleep
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi.testclient import TestClient

from docchecker.api import main
from docchecker.api.main import app
from docchecker.domain.document import UploadedDocumentRecord
from docchecker.domain.enums import SourceType, TaskStatus
from docchecker.domain.requirements import RequirementDocument
from docchecker.domain.rules import RuleSet
from docchecker.services.file_storage import LocalFileStorage
from docchecker.services.state_store import SqliteStateStore
from docchecker.services.task_queue import BackgroundJobEnqueueError
from docchecker.services.word_document_preparer import PreparedWordDocument


def _create_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def _fake_prepare_doc_upload(path: Path, **kwargs) -> PreparedWordDocument:
    normalized_path = path.with_suffix(".docx")
    _create_docx(normalized_path, "正文内容")
    return PreparedWordDocument(
        original_path=path,
        normalized_path=normalized_path,
        original_format=path.suffix.lower().lstrip("."),
        normalized_format="docx",
    )


def _wait_for_json(client: TestClient, path: str, predicate, *, attempts: int = 100) -> dict:
    last_payload: dict | None = None
    for _ in range(attempts):
        response = client.get(path)
        assert response.status_code == 200
        payload = response.json()
        last_payload = payload
        if predicate(payload):
            return payload
        sleep(0.05)
    raise AssertionError(f"Timed out waiting for {path}: {last_payload}")


@pytest.fixture(autouse=True)
def isolate_api_runtime(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    store = SqliteStateStore(tmp_path / "docchecker.sqlite3")
    store.initialize()
    monkeypatch.setattr(main, "state_store", store)
    monkeypatch.setattr(main, "storage", LocalFileStorage(tmp_path / "storage"))


def test_manual_ruleset_check_flow(tmp_path: Path) -> None:
    docx_path = tmp_path / "paper.docx"
    _create_docx(docx_path, "正文内容")
    client = TestClient(app)

    with docx_path.open("rb") as file:
        document_response = client.post(
            "/api/documents",
            files={
                "file": (
                    "paper.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert document_response.status_code == 200
    document_id = document_response.json()["document_id"]

    draft_response = client.post(
        "/api/draft-rulesets",
        json={
            "document_id": document_id,
            "source_type": "manual",
            "manual_text": "正文宋体小四，1.5倍行距，页边距上下2.5cm。",
        },
    )
    assert draft_response.status_code == 200
    draft = draft_response.json()
    assert draft["rules"]
    assert draft["extraction_summary"]["structured_rules"] == len(draft["rules"])
    assert "unsupported_requirements" in draft

    publish_response = client.post(f"/api/draft-rulesets/{draft['id']}/publish")
    assert publish_response.status_code == 200
    ruleset_id = publish_response.json()["id"]

    task_response = client.post(
        "/api/check-tasks",
        json={"document_id": document_id, "ruleset_id": ruleset_id},
    )
    assert task_response.status_code == 200
    task = task_response.json()
    assert task["status"] == "pending"
    assert task["document_filename"] == "paper.docx"
    assert task["ruleset_name"] == "候选规则集"
    task = _wait_for_json(
        client,
        f"/api/check-tasks/{task['id']}",
        lambda payload: payload["status"] == "succeeded",
    )
    assert task["status"] == "succeeded"
    assert task["report_id"]

    history_response = client.get("/api/check-tasks")
    assert history_response.status_code == 200
    assert any(item["id"] == task["id"] for item in history_response.json())
    paged_history_response = client.get("/api/check-tasks?limit=1&offset=0")
    assert paged_history_response.status_code == 200
    assert len(paged_history_response.json()) == 1

    restarted_store = SqliteStateStore(main.state_store.path)
    restarted_store.initialize()
    main.state_store = restarted_store

    restarted_history_response = client.get("/api/check-tasks")
    assert restarted_history_response.status_code == 200
    assert any(item["id"] == task["id"] for item in restarted_history_response.json())

    restarted_task_response = client.get(f"/api/check-tasks/{task['id']}")
    assert restarted_task_response.status_code == 200
    assert restarted_task_response.json()["report_id"] == task["report_id"]

    report_response = client.get(f"/api/reports/{task['report_id']}")
    assert report_response.status_code == 200
    assert report_response.json()["id"] == task["report_id"]
    assert report_response.json()["document_filename"] == "paper.docx"
    assert report_response.json()["ruleset_name"] == "候选规则集"

    draft_response = client.get(f"/api/draft-rulesets/{draft['id']}")
    assert draft_response.status_code == 200
    assert draft_response.json()["status"] == "published"

    rulesets_response = client.get("/api/rulesets")
    assert rulesets_response.status_code == 200
    assert any(item["id"] == ruleset_id for item in rulesets_response.json())

    rename_response = client.patch(
        f"/api/rulesets/{ruleset_id}",
        json={"name": "学校论文格式模板"},
    )
    assert rename_response.status_code == 200
    renamed_ruleset = rename_response.json()
    assert renamed_ruleset["name"] == "学校论文格式模板"
    assert renamed_ruleset["rules"]

    main.state_store = SqliteStateStore(main.state_store.path)
    main.state_store.initialize()
    renamed_list_response = client.get("/api/rulesets")
    assert renamed_list_response.status_code == 200
    assert any(
        item["id"] == ruleset_id and item["name"] == "学校论文格式模板"
        for item in renamed_list_response.json()
    )

    delete_task_response = client.delete(f"/api/check-tasks/{task['id']}")
    assert delete_task_response.status_code == 200
    assert delete_task_response.json() == {"id": task["id"], "deleted": True}
    assert client.get(f"/api/check-tasks/{task['id']}").status_code == 404
    assert client.get(f"/api/reports/{task['report_id']}").status_code == 404
    assert not main.storage.report_path(task["report_id"]).exists()

    delete_ruleset_response = client.delete(f"/api/rulesets/{ruleset_id}")
    assert delete_ruleset_response.status_code == 200
    assert delete_ruleset_response.json() == {"id": ruleset_id, "deleted": True}
    assert not any(item["id"] == ruleset_id for item in client.get("/api/rulesets").json())
    paged_ruleset_response = client.get("/api/rulesets?limit=1&offset=0")
    assert paged_ruleset_response.status_code == 200
    assert len(paged_ruleset_response.json()) <= 1
    assert client.delete(f"/api/rulesets/{ruleset_id}").status_code == 404
    assert client.delete(f"/api/check-tasks/{task['id']}").status_code == 404


def test_requirement_document_draft_is_created_before_background_extraction(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    docx_path = tmp_path / "paper.docx"
    rules_path = tmp_path / "rules.docx"
    _create_docx(docx_path, "正文内容")
    _create_docx(rules_path, "正文宋体小四。")
    client = TestClient(app)

    def fake_extract_requirement_document(requirement_model, *, source_type):
        return main.extract_rules_from_text("正文宋体小四。", source_type=source_type)

    scheduled_jobs: list[tuple] = []

    def capture_background_job(function, *args) -> None:
        scheduled_jobs.append((function, args))

    monkeypatch.setattr(
        main,
        "extract_rules_from_requirement_document",
        fake_extract_requirement_document,
    )
    monkeypatch.setattr(main, "_start_background_job", capture_background_job)

    with docx_path.open("rb") as file:
        document_response = client.post(
            "/api/documents",
            files={
                "file": (
                    "paper.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    document_id = document_response.json()["document_id"]

    with rules_path.open("rb") as file:
        requirement_response = client.post(
            "/api/requirement-documents",
            files={
                "file": (
                    "rules.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    requirement_id = requirement_response.json()["id"]

    draft_response = client.post(
        "/api/draft-rulesets",
        json={
            "document_id": document_id,
            "source_type": "requirement_doc",
            "requirement_document_id": requirement_id,
        },
    )

    assert draft_response.status_code == 200
    draft = draft_response.json()
    assert draft["status"] == "processing"
    assert draft["rules"] == []
    assert len(scheduled_jobs) == 1

    function, args = scheduled_jobs[0]
    function(*args)

    persisted = _wait_for_json(
        client,
        f"/api/draft-rulesets/{draft['id']}",
        lambda payload: payload["status"] == "draft",
    )
    assert persisted["status"] == "draft"
    assert persisted["rules"]
    assert persisted["error"] is None


def test_check_task_uses_rq_enqueue_without_inline_execution(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document_path = tmp_path / "paper.docx"
    _create_docx(document_path, "正文内容")
    main.state_store.save_document(
        UploadedDocumentRecord(
            id="doc_rq",
            filename="paper.docx",
            path=str(document_path),
            original_path=str(document_path),
            original_format="docx",
            normalized_format="docx",
            size_bytes=document_path.stat().st_size,
        )
    )
    main.state_store.save_ruleset(
        RuleSet(
            id="ruleset_rq",
            name="RQ 模板",
            source_type=SourceType.manual,
            version="1.0.0",
            rules=[],
            created_at="2026-05-12T00:00:00+00:00",
        )
    )
    enqueued_jobs: list[tuple] = []

    def capture_start_background_job(settings, function, *args):
        enqueued_jobs.append((settings.task_execution_mode, function, args))
        return "rq_job_1"

    monkeypatch.setattr(main.settings, "task_execution_mode", "rq")
    monkeypatch.setattr(main, "start_background_job", capture_start_background_job)
    client = TestClient(app)

    task_response = client.post(
        "/api/check-tasks",
        json={"document_id": "doc_rq", "ruleset_id": "ruleset_rq"},
    )

    assert task_response.status_code == 200
    task = task_response.json()
    assert task["status"] == "pending"
    assert enqueued_jobs == [("rq", main._execute_check_task, (task["id"],))]
    persisted_task = main.state_store.get_check_task(task["id"])
    assert persisted_task is not None
    assert persisted_task.status == TaskStatus.pending
    assert persisted_task.report_id is None


def test_check_task_marks_failed_when_rq_enqueue_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document_path = tmp_path / "paper.docx"
    _create_docx(document_path, "正文内容")
    main.state_store.save_document(
        UploadedDocumentRecord(
            id="doc_rq_fail",
            filename="paper.docx",
            path=str(document_path),
            original_path=str(document_path),
            original_format="docx",
            normalized_format="docx",
            size_bytes=document_path.stat().st_size,
        )
    )
    main.state_store.save_ruleset(
        RuleSet(
            id="ruleset_rq_fail",
            name="RQ 模板",
            source_type=SourceType.manual,
            version="1.0.0",
            rules=[],
            created_at="2026-05-12T00:00:00+00:00",
        )
    )

    def fail_start_background_job(settings, function, *args):
        raise BackgroundJobEnqueueError("Redis 不可用")

    monkeypatch.setattr(main.settings, "task_execution_mode", "rq")
    monkeypatch.setattr(main, "start_background_job", fail_start_background_job)
    client = TestClient(app, raise_server_exceptions=False)

    task_response = client.post(
        "/api/check-tasks",
        json={"document_id": "doc_rq_fail", "ruleset_id": "ruleset_rq_fail"},
    )

    assert task_response.status_code == 503
    assert task_response.json()["detail"] == "Redis 不可用"
    persisted_tasks = main.state_store.list_check_tasks()
    assert len(persisted_tasks) == 1
    assert persisted_tasks[0].status == TaskStatus.failed
    assert persisted_tasks[0].error == "Redis 不可用"


def test_requirement_draft_marks_failed_when_rq_enqueue_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document_path = tmp_path / "paper.docx"
    requirement_path = tmp_path / "rules.docx"
    _create_docx(document_path, "正文内容")
    _create_docx(requirement_path, "正文宋体小四。")
    main.state_store.save_document(
        UploadedDocumentRecord(
            id="doc_draft_rq_fail",
            filename="paper.docx",
            path=str(document_path),
            original_path=str(document_path),
            original_format="docx",
            normalized_format="docx",
            size_bytes=document_path.stat().st_size,
        )
    )
    main.state_store.save_requirement_document(
        RequirementDocument(
            id="req_draft_rq_fail",
            filename="rules.docx",
            path=str(requirement_path),
            size_bytes=requirement_path.stat().st_size,
            extracted_text="正文宋体小四。",
            original_format="docx",
            normalized_format="docx",
            created_at="2026-05-12T00:00:00+00:00",
        )
    )

    def fail_start_background_job(settings, function, *args):
        raise BackgroundJobEnqueueError("Redis 不可用")

    monkeypatch.setattr(main.settings, "task_execution_mode", "rq")
    monkeypatch.setattr(main, "start_background_job", fail_start_background_job)
    monkeypatch.setattr(main, "uuid4", lambda: SimpleNamespace(hex="draft_rq_fail"))
    client = TestClient(app, raise_server_exceptions=False)

    draft_response = client.post(
        "/api/draft-rulesets",
        json={
            "document_id": "doc_draft_rq_fail",
            "source_type": "requirement_doc",
            "requirement_document_id": "req_draft_rq_fail",
        },
    )

    assert draft_response.status_code == 503
    assert draft_response.json()["detail"] == "Redis 不可用"
    persisted_draft = main.state_store.get_draft_ruleset("draft_draft_rq_fail")
    assert persisted_draft is not None
    assert persisted_draft.status.value == "failed"
    assert persisted_draft.error == "Redis 不可用"


def test_upload_document_accepts_doc_after_conversion(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(main, "prepare_word_document", _fake_prepare_doc_upload)
    client = TestClient(app)

    document_response = client.post(
        "/api/documents",
        files={"file": ("paper.doc", b"legacy word content", "application/msword")},
    )

    assert document_response.status_code == 200
    payload = document_response.json()
    assert payload["filename"] == "paper.doc"
    assert payload["original_format"] == "doc"
    assert payload["normalized_format"] == "docx"
    document = main.state_store.get_document(payload["document_id"])
    assert document is not None
    assert document.path.endswith(".docx")


def test_upload_requirement_document_accepts_doc_after_conversion(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(main, "prepare_word_document", _fake_prepare_doc_upload)
    client = TestClient(app)

    requirement_response = client.post(
        "/api/requirement-documents",
        files={"file": ("rules.doc", b"legacy word content", "application/msword")},
    )

    assert requirement_response.status_code == 200
    payload = requirement_response.json()
    assert payload["filename"] == "rules.doc"
    assert payload["path"].endswith(".docx")
    assert payload["original_format"] == "doc"
    assert payload["normalized_format"] == "docx"
    assert "正文内容" in payload["extracted_text"]


def test_upload_document_cleans_file_when_persistence_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FailingStore:
        def save_document(self, document) -> None:
            raise RuntimeError("database unavailable")

    docx_path = tmp_path / "paper.docx"
    _create_docx(docx_path, "正文内容")
    monkeypatch.setattr(main, "state_store", FailingStore())
    client = TestClient(app, raise_server_exceptions=False)

    with docx_path.open("rb") as file:
        document_response = client.post(
            "/api/documents",
            files={
                "file": (
                    "paper.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert document_response.status_code == 500
    assert list(main.storage.documents_dir.iterdir()) == []


def test_upload_document_rejects_oversized_file_during_stream(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def fail_prepare_word_document(*args, **kwargs):
        raise AssertionError("prepare_word_document should not run for oversized uploads")

    monkeypatch.setattr(main.settings, "max_document_size_bytes", 4)
    monkeypatch.setattr(main, "prepare_word_document", fail_prepare_word_document)
    client = TestClient(app, raise_server_exceptions=False)

    document_response = client.post(
        "/api/documents",
        files={
            "file": (
                "paper.docx",
                BytesIO(b"abcde"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert document_response.status_code == 400
    assert "超过限制" in document_response.json()["detail"]
    assert list(main.storage.documents_dir.iterdir()) == []


def test_check_task_marks_failed_and_cleans_report_when_final_persistence_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    docx_path = tmp_path / "paper.docx"
    _create_docx(docx_path, "正文内容")
    client = TestClient(app)

    with docx_path.open("rb") as file:
        document_response = client.post(
            "/api/documents",
            files={
                "file": (
                    "paper.docx",
                    file,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert document_response.status_code == 200
    document_id = document_response.json()["document_id"]

    ruleset_id = "ruleset_finalize_fail"
    ruleset_response = client.post(
        "/api/rulesets",
        json={
            "id": ruleset_id,
            "name": "空规则集",
            "source_type": "manual",
            "version": "1.0.0",
            "created_at": "2026-04-26T00:00:00+08:00",
            "rules": [],
        },
    )
    assert ruleset_response.status_code == 200

    def fail_save_many(records) -> None:
        raise RuntimeError("database unavailable")

    monkeypatch.setattr(main.state_store, "save_many", fail_save_many)
    task_response = client.post(
        "/api/check-tasks",
        json={"document_id": document_id, "ruleset_id": ruleset_id},
    )

    assert task_response.status_code == 200
    task = task_response.json()
    assert task["status"] == "pending"
    assert list(main.storage.reports_dir.iterdir()) == []
    persisted_task = _wait_for_json(
        client,
        f"/api/check-tasks/{task['id']}",
        lambda payload: payload["status"] == "failed",
    )
    assert persisted_task["status"] == "failed"
    assert persisted_task["error"] == "database unavailable"
