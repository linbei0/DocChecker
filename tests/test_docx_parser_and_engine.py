from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from docchecker.checkers.engine import CheckEngine
from docchecker.domain.enums import RuleCategory, Severity, SourceType
from docchecker.domain.findings import CheckReport
from docchecker.domain.rules import FormatRule, RuleSet, RuleSource, RuleTarget
from docchecker.parsing.docx_parser import parse_docx
from docchecker.reports.markdown import render_markdown_report
from docchecker.services.docx_validator import validate_docx_path


def _create_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    paragraph = document.add_paragraph("正文内容")
    run = paragraph.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(12)
    document.save(path)


def test_parse_docx_reads_sections_and_paragraphs(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    _create_docx(path)

    validate_docx_path(path, max_size_bytes=1024 * 1024)
    model = parse_docx(path, document_id="doc_1", source_filename="sample.docx")

    assert model.sections[0].margin_top_cm == pytest.approx(2.5, abs=0.01)
    assert model.paragraphs[0].text == "正文内容"
    assert model.paragraphs[0].font_size_pt == 12


def test_parse_docx_builds_document_facts(tmp_path: Path) -> None:
    path = tmp_path / "facts.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "学校论文"
    document.add_paragraph("正文内容")
    document.add_paragraph("表1.1 基本要求")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "要求"
    document.add_paragraph("表后说明")
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="facts.docx")

    assert any(fact.text == "学校论文" for fact in model.facts.headers_footers)
    header = next(fact for fact in model.facts.headers_footers if fact.text == "学校论文")
    assert header.paragraphs[0].text == "学校论文"
    assert header.paragraphs[0].effective_format_sources
    assert model.facts.tables[0].row_count == 1
    assert model.facts.tables[0].cells[1].text == "要求"
    assert model.facts.tables[0].caption_text == "表1.1 基本要求"
    assert model.facts.tables[0].caption_position == "before"
    assert model.facts.tables[0].preceding_paragraph_index == 1
    assert model.facts.tables[0].following_paragraph_index == 4
    assert model.paragraphs[2].text == "项目"
    assert model.paragraphs[2].table_index == 0
    assert model.paragraphs[2].row_index == 0
    assert model.paragraphs[2].column_index == 0
    assert model.paragraphs[2].cell_paragraph_index == 0
    assert "word/document.xml" in model.facts.xml_parts


def test_parse_docx_resolves_style_inherited_paragraph_format(tmp_path: Path) -> None:
    path = tmp_path / "styled.docx"
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(6)
    document.add_paragraph("结论")
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="styled.docx")
    paragraph = model.paragraphs[0]

    assert paragraph.font_family == "宋体"
    assert paragraph.font_size_pt == 12
    assert paragraph.alignment == "left"
    assert paragraph.space_after_pt == 6
    assert paragraph.raw["effective_format_sources"]["font_size_pt"] == "paragraph_style:Normal"
    assert paragraph.raw["effective_format_sources"]["space_after_pt"] == "paragraph_style:Normal"
    fact = next(item for item in model.facts.effective_formats if item.owner_id == "paragraph:0")
    assert fact.values["font_size_pt"] == 12
    assert fact.sources["font_size_pt"] == "paragraph_style:Normal"


def test_property_and_ooxml_checkers_use_document_facts(tmp_path: Path) -> None:
    path = tmp_path / "property.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "论文"
    document.add_paragraph("正文内容")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="property.docx")
    rules = [
        FormatRule(
            id="fact_header_contains_school",
            category=RuleCategory.header_footer,
            target=RuleTarget(scope="document.header_footer"),
            expectation={
                "$facts": [
                    {
                        "path": "facts.headers_footers.text",
                        "operator": "contains",
                        "value": "学校",
                    }
                ]
            },
            severity=Severity.minor,
            source=RuleSource(
                type=SourceType.manual,
                excerpt="页眉应包含学校。",
            ),
        ),
        FormatRule(
            id="ooxml_requires_body",
            category=RuleCategory.structure,
            target=RuleTarget(scope="document.structure"),
            expectation={
                "$xpath": [
                    {
                        "part": "word/document.xml",
                        "expression": "boolean(/w:document/w:body)",
                    }
                ]
            },
            severity=Severity.major,
            source=RuleSource(
                type=SourceType.manual,
                excerpt="文档必须包含 body。",
            ),
        ),
    ]

    findings = CheckEngine().run(model, rules)

    assert any(finding.checker_id == "property" for finding in findings)
    assert not any(finding.rule_id == "ooxml_requires_body" for finding in findings)


def test_unified_rule_dsl_routes_to_property_and_ooxml_backends(tmp_path: Path) -> None:
    path = tmp_path / "dsl.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "论文"
    document.add_paragraph("正文内容")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="dsl.docx")
    rules = [
        FormatRule(
            id="dsl_header_contains_school",
            category=RuleCategory.header_footer,
            target=RuleTarget(scope="document.header_footer"),
            expectation={
                "$dsl": [
                    {
                        "backend": "facts",
                        "path": "facts.headers_footers.text",
                        "operator": "contains",
                        "value": "学校",
                    }
                ]
            },
            severity=Severity.minor,
            source=RuleSource(type=SourceType.manual, excerpt="页眉应包含学校。"),
        ),
        FormatRule(
            id="dsl_ooxml_requires_body",
            category=RuleCategory.structure,
            target=RuleTarget(scope="document.structure"),
            expectation={
                "$dsl": [
                    {
                        "backend": "ooxml",
                        "operator": "xpath",
                        "part": "word/document.xml",
                        "expression": "boolean(/w:document/w:body)",
                    }
                ]
            },
            severity=Severity.major,
            source=RuleSource(type=SourceType.manual, excerpt="文档必须包含 body。"),
        ),
    ]

    findings = CheckEngine().run(model, rules)

    assert any(finding.rule_id == "dsl_header_contains_school" for finding in findings)
    assert not any(finding.rule_id == "dsl_ooxml_requires_body" for finding in findings)


def test_parse_docx_reads_simple_field_facts(tmp_path: Path) -> None:
    path = tmp_path / "field.docx"
    document = Document()
    paragraph = document.add_paragraph("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    paragraph.add_run(" 页")
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="field.docx")

    assert any(fact.field_type == "PAGE" for fact in model.facts.fields)
    assert model.facts.fields[0].part_name == "word/document.xml"


def test_header_footer_scope_checks_footer_page_number_xml(tmp_path: Path) -> None:
    path = tmp_path / "footer-page.docx"
    document = Document()
    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.text = ""
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer_paragraph._p.append(field)
    document.add_paragraph("正文内容")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="footer-page.docx")
    rule = FormatRule(
        id="header_footer_basic",
        category=RuleCategory.header_footer,
        target=RuleTarget(scope="document.header_footer"),
        expectation={"requiresPageNumber": True},
        severity=Severity.minor,
        source=RuleSource(type=SourceType.manual, excerpt="页眉页脚应包含页码。"),
    )

    findings = CheckEngine().run(model, [rule])

    assert not any(finding.rule_id == "header_footer_basic" for finding in findings)


def test_parse_docx_builds_phase4_high_frequency_facts(tmp_path: Path) -> None:
    path = tmp_path / "phase4-facts.docx"
    document = Document()
    document.add_paragraph("摘 要")
    document.add_paragraph("本文研究论文格式检查系统。")
    document.add_paragraph("关键词：论文；格式检查")
    document.add_paragraph("目 录")
    document.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    toc_field_paragraph = document.add_paragraph("1 绪论 1", style="toc 1")
    toc_field = OxmlElement("w:fldSimple")
    toc_field.set(qn("w:instr"), r"TOC \o \"1-3\"")
    toc_field_paragraph._p.append(toc_field)
    document.add_paragraph("参考文献")
    document.add_paragraph("[1] 张三. 论文格式研究.")
    document.add_paragraph("表1.1 基本要求")
    document.add_table(rows=1, cols=1)
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="phase4-facts.docx")

    assert model.facts.toc.has_title is True
    assert model.facts.toc.has_field is True
    assert model.facts.toc.entry_count == 1
    assert model.facts.references.has_section is True
    assert model.facts.references.entry_count == 1
    assert model.facts.references.numbering_continuous is True
    assert model.facts.abstracts[0].language == "zh"
    assert model.facts.abstracts[0].has_keywords is True
    assert model.facts.abstracts[0].keyword_count == 2
    assert model.facts.captions[0].kind == "table"
    assert model.facts.captions[0].position == "before"


def test_parse_docx_resolves_east_asia_font_from_style_xml(tmp_path: Path) -> None:
    path = tmp_path / "heading-font.docx"
    document = Document()
    style = document.styles["Heading 1"]
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "黑体")
    style.font.size = Pt(16)
    document.add_paragraph("绪论", style="Heading 1")
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="heading-font.docx")
    paragraph = model.paragraphs[0]

    assert paragraph.font_family == "黑体"
    assert paragraph.font_size_pt == 16


def test_parse_docx_tracks_mixed_script_fonts_separately(tmp_path: Path) -> None:
    path = tmp_path / "mixed-font.docx"
    document = Document()
    paragraph = document.add_paragraph()
    chinese = paragraph.add_run("关键词：校园社团管理")
    chinese._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    english = paragraph.add_run(" Layui SpringBoot")
    english.font.name = "Times New Roman"
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="mixed-font.docx")
    paragraph_node = model.paragraphs[0]

    assert paragraph_node.font_family is None
    assert paragraph_node.font_family_east_asia == "宋体"
    assert paragraph_node.font_family_ascii == "Times New Roman"
    assert [run.script for run in paragraph_node.runs] == ["east_asia", "ascii"]
    assert paragraph_node.runs[0].font_family_east_asia == "宋体"
    assert paragraph_node.runs[1].font_family_ascii == "Times New Roman"


def test_font_checker_ignores_keyword_label_when_checking_keyword_content_font(
    tmp_path: Path,
) -> None:
    path = tmp_path / "mixed-font-check.docx"
    document = Document()
    paragraph = document.add_paragraph()
    label = paragraph.add_run("关键词：")
    label._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    content = paragraph.add_run("校园社团管理")
    content._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    english = paragraph.add_run(" Layui")
    english.font.name = "Times New Roman"
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="mixed-font-check.docx")
    ruleset = RuleSet(
        id="ruleset_mixed_font",
        name="混排字体",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="keyword_cn_font",
                category=RuleCategory.font,
                target=RuleTarget(scope="keywords.paragraph", selector="关键词"),
                expectation={"fontFamilyEastAsia": "宋体"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="关键词宋体"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert findings == []


def test_font_checker_checks_keyword_english_runs_separately(tmp_path: Path) -> None:
    path = tmp_path / "keyword-english-font.docx"
    document = Document()
    paragraph = document.add_paragraph()
    label = paragraph.add_run("关键词：")
    label._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    content = paragraph.add_run("校园社团管理")
    content._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    english = paragraph.add_run(" Layui")
    english.font.name = "Times New Roman"
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="keyword-english-font.docx")
    ruleset = RuleSet(
        id="ruleset_keyword_english_font",
        name="关键词英文字体",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="keyword_en_font",
                category=RuleCategory.font,
                target=RuleTarget(scope="keywords.paragraph", selector="关键词"),
                expectation={"fontFamilyEastAsia": "Times New Roman"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="关键词英文 Times New Roman"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert findings == []


def test_font_checker_does_not_fail_keyword_chinese_when_only_latin_fallback_exists(
    tmp_path: Path,
) -> None:
    path = tmp_path / "keyword-latin-fallback.docx"
    document = Document()
    paragraph = document.add_paragraph()
    label = paragraph.add_run("关键词：")
    label._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    content = paragraph.add_run("校园社团管理")
    content.font.name = "Calibri"
    english = paragraph.add_run(" Layui")
    english.font.name = "Times New Roman"
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="keyword-latin-fallback.docx")
    ruleset = RuleSet(
        id="ruleset_keyword_cn_fallback",
        name="关键词中文字体",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="keyword_cn_font",
                category=RuleCategory.font,
                target=RuleTarget(scope="keywords.paragraph", selector="关键词"),
                expectation={"fontFamilyEastAsia": "宋体"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="关键词中文宋体"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert findings == []


def test_parse_docx_tracks_nearest_heading_as_logical_section(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.add_paragraph("研究背景")
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="sample.docx")

    assert model.paragraphs[0].raw["section_name"] == "绪论"
    assert model.paragraphs[1].raw["section_name"] == "绪论"
    assert model.logical_sections[0].role == "body"
    assert model.logical_sections[0].start_paragraph_index == 0


def test_engine_reports_font_mismatch(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    _create_docx(path)
    model = parse_docx(path, document_id="doc_1", source_filename="sample.docx")
    ruleset = RuleSet(
        id="ruleset_test",
        name="测试规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="body_font_family",
                category=RuleCategory.font,
                target=RuleTarget(scope="body.paragraph"),
                expectation={"fontFamilyEastAsia": "宋体"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="正文宋体"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert len(findings) == 1
    assert findings[0].rule_id == "body_font_family"
    assert findings[0].actual == {"fontFamilyEastAsia": "Arial"}
    assert findings[0].excerpt == "正文内容"
    assert findings[0].location.paragraph_number == 1
    assert findings[0].location.display_path == "第 1 段"
    assert findings[0].context["style_name"] == "Normal"
    assert findings[0].context["field_label"] == "中文字体"


def test_body_rules_ignore_table_cells_unless_scope_is_explicit(tmp_path: Path) -> None:
    path = tmp_path / "table-body.docx"
    document = Document()
    paragraph = document.add_paragraph("正文")
    paragraph.runs[0].font.name = "Arial"
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格正文"
    run = table.cell(0, 0).paragraphs[0].runs[0]
    run.font.name = "Arial"
    document.save(path)

    model = parse_docx(path, document_id="doc_1", source_filename="table-body.docx")
    rules = [
        FormatRule(
            id="body_font",
            category=RuleCategory.font,
            target=RuleTarget(scope="body.paragraph"),
            expectation={"fontFamilyEastAsia": "宋体"},
            severity=Severity.major,
            source=RuleSource(type=SourceType.manual, excerpt="正文宋体"),
        ),
        FormatRule(
            id="table_cell_font",
            category=RuleCategory.font,
            target=RuleTarget(scope="table_cell"),
            expectation={"fontFamilyEastAsia": "宋体"},
            severity=Severity.major,
            source=RuleSource(type=SourceType.manual, excerpt="表格正文宋体"),
        ),
    ]

    findings = CheckEngine().run(model, rules)

    assert any(
        finding.rule_id == "body_font" and finding.excerpt == "正文"
        for finding in findings
    )
    assert not any(
        finding.rule_id == "body_font" and finding.excerpt == "表格正文"
        for finding in findings
    )
    table_finding = next(finding for finding in findings if finding.excerpt == "表格正文")
    assert table_finding.rule_id == "table_cell_font"
    assert table_finding.location.display_path == (
        "正文 / 表 1 第 1 行第 1 列 / 单元格第 1 段 / 第 2 段"
    )
    assert table_finding.location.table_index == 0
    assert table_finding.context["column_index"] == 0


def test_engine_logs_checker_failure(caplog: pytest.LogCaptureFixture) -> None:
    class FailingChecker:
        checker_id = "failing"

        def check(self, document, rules):
            raise RuntimeError("checker boom")

    caplog.set_level("ERROR", logger="docchecker.checkers.engine")

    findings = CheckEngine(checkers=[FailingChecker()]).run(object(), [])

    assert len(findings) == 1
    assert findings[0].rule_id == "__checker_execution__"
    assert findings[0].actual == {"error": "checker boom"}
    assert any(
        record.exc_info and "检查器 failing 执行失败" in record.message
        for record in caplog.records
    )


def test_heading_rules_do_not_match_body_text_mentions(tmp_path: Path) -> None:
    path = tmp_path / "heading-selector.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.add_paragraph("第一章：绪论。本章介绍研究背景。")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="heading-selector.docx")
    ruleset = RuleSet(
        id="ruleset_heading",
        name="标题规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="heading_alignment",
                category=RuleCategory.heading,
                target=RuleTarget(scope="heading.level1", selector="绪论"),
                expectation={"alignment": "left"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="一级标题左对齐"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert len(findings) == 1
    assert findings[0].excerpt == "绪论"


def test_heading_rules_match_exact_heading_level(tmp_path: Path) -> None:
    path = tmp_path / "heading-levels.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="heading-levels.docx")
    ruleset = RuleSet(
        id="ruleset_heading_levels",
        name="标题级别规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="heading3_alignment",
                category=RuleCategory.paragraph,
                target=RuleTarget(scope="heading.3", selector="Heading 3"),
                expectation={"alignment": "left"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="三级标题左对齐"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert findings == []


def test_missing_paragraph_defaults_are_reported_as_values(tmp_path: Path) -> None:
    path = tmp_path / "paragraph-defaults.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="paragraph-defaults.docx")
    ruleset = RuleSet(
        id="ruleset_defaults",
        name="段落默认值",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="heading1_indent",
                category=RuleCategory.paragraph,
                target=RuleTarget(scope="heading.1", selector="Heading 1"),
                expectation={"firstLineIndentCm": 0.74},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="一级标题首行缩进2字符"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert len(findings) == 1
    assert findings[0].actual == {"firstLineIndentCm": 0.0}
    assert findings[0].certainty.value == "certain"


def test_markdown_report_includes_fragment_context(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    _create_docx(path)
    model = parse_docx(path, document_id="doc_1", source_filename="sample.docx")
    ruleset = RuleSet(
        id="ruleset_test",
        name="测试规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="body_font_family",
                category=RuleCategory.font,
                target=RuleTarget(scope="body.paragraph"),
                expectation={"fontFamilyEastAsia": "宋体"},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="正文宋体"),
                confidence=1,
            )
        ],
    )
    findings = CheckEngine().run(model, ruleset.rules)
    report = CheckReport(
        id="report_test",
        document_id="doc_1",
        ruleset_id=ruleset.id,
        checker_version="0.1.0",
        generated_at="2026-04-26T00:00:00+08:00",
        findings=findings,
    )

    content = render_markdown_report(report)

    assert "## 问题片段" in content
    assert "第 1 段" in content
    assert "正文内容" in content
    assert "期望值：fontFamilyEastAsia=宋体" in content
    assert "实际值：fontFamilyEastAsia=Arial" in content


def test_semantic_checkers_report_structure_and_reference_findings(tmp_path: Path) -> None:
    path = tmp_path / "paper.docx"
    document = Document()
    document.add_paragraph("摘要")
    document.add_paragraph("正文内容")
    document.add_paragraph("[2] 不连续编号参考文献")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="paper.docx")
    ruleset = RuleSet(
        id="ruleset_semantic",
        name="语义规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="structure_required_sections",
                category=RuleCategory.structure,
                target=RuleTarget(scope="document.structure", selector="论文结构"),
                expectation={"requiredSections": ["摘要", "目录", "正文"]},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="论文应包含摘要、目录、正文"),
                confidence=1,
            ),
            FormatRule(
                id="reference_basic_entries",
                category=RuleCategory.reference,
                target=RuleTarget(scope="document.references", selector="参考文献"),
                expectation={"requiresReferences": True, "numbering": "bracketed"},
                severity=Severity.minor,
                source=RuleSource(type=SourceType.manual, excerpt="参考文献按编号排列"),
                confidence=1,
            ),
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)
    rule_ids = {finding.rule_id for finding in findings}

    assert "structure_required_sections" in rule_ids
    assert "reference_basic_entries" in rule_ids


def test_phase4_checkers_use_high_frequency_facts(tmp_path: Path) -> None:
    path = tmp_path / "phase4-checkers.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "论文"
    document.add_paragraph("摘 要")
    document.add_paragraph("短摘要")
    document.add_paragraph("目 录")
    document.add_paragraph("参考文献")
    document.add_paragraph("[2] 编号不连续参考文献")
    document.add_table(rows=1, cols=1)
    document.add_paragraph("表1.1 基本要求")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="phase4-checkers.docx")
    rules = [
        FormatRule(
            id="header_font",
            category=RuleCategory.header_footer,
            target=RuleTarget(scope="header.default"),
            expectation={"fontSizePt": 10.5},
            severity=Severity.minor,
            source=RuleSource(type=SourceType.manual, excerpt="页眉五号字"),
        ),
        FormatRule(
            id="toc_field",
            category=RuleCategory.toc,
            target=RuleTarget(scope="document.toc"),
            expectation={"requiresToc": True, "requiresTocField": True},
            severity=Severity.major,
            source=RuleSource(type=SourceType.manual, excerpt="目录自动生成"),
        ),
        FormatRule(
            id="caption_position",
            category=RuleCategory.caption,
            target=RuleTarget(scope="document.caption"),
            expectation={"requiresTableCaption": True, "tableCaptionPosition": "before"},
            severity=Severity.minor,
            source=RuleSource(type=SourceType.manual, excerpt="表题置于表上"),
        ),
        FormatRule(
            id="reference_continuous",
            category=RuleCategory.reference,
            target=RuleTarget(scope="document.references"),
            expectation={
                "requiresSection": True,
                "requiresReferences": True,
                "numberingContinuous": True,
            },
            severity=Severity.major,
            source=RuleSource(type=SourceType.manual, excerpt="参考文献连续编号"),
        ),
        FormatRule(
            id="abstract_requirements",
            category=RuleCategory.abstract,
            target=RuleTarget(scope="document.abstract"),
            expectation={
                "requiresChineseAbstract": True,
                "requiresKeywords": True,
                "minWordCount": 10,
            },
            severity=Severity.major,
            source=RuleSource(type=SourceType.manual, excerpt="摘要和关键词要求"),
        ),
    ]

    findings = CheckEngine().run(model, rules)
    rule_ids = {finding.rule_id for finding in findings}

    assert "header_font" in rule_ids
    assert "toc_field" in rule_ids
    assert "caption_position" in rule_ids
    assert "reference_continuous" in rule_ids
    assert "abstract_requirements" in rule_ids


def test_structure_checker_matches_common_section_aliases(tmp_path: Path) -> None:
    path = tmp_path / "paper-alias.docx"
    document = Document()
    document.add_paragraph("摘 要：系统设计与实现。")
    document.add_paragraph("关键词：校园社团管理；Layui")
    document.add_paragraph("目  录")
    document.add_heading("绪论", level=1)
    document.add_paragraph("正文内容")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="paper-alias.docx")
    ruleset = RuleSet(
        id="ruleset_structure_alias",
        name="结构规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="structure_required_sections",
                category=RuleCategory.structure,
                target=RuleTarget(scope="document.structure", selector="论文结构"),
                expectation={
                    "requiredSections": ["中文摘要", "中文关键词", "目录", "正文"]
                },
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="论文结构要求"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert findings == []


def test_structure_checker_ignores_toc_entries_for_order(tmp_path: Path) -> None:
    path = tmp_path / "paper-toc-order.docx"
    document = Document()
    document.add_paragraph("摘 要：系统设计与实现。")
    document.add_paragraph("关键词：校园社团管理；Layui")
    document.add_paragraph("目  录")
    document.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("致  谢\t55", style="toc 1")
    document.add_heading("绪论", level=1)
    document.add_paragraph("正文内容")
    document.add_paragraph("致  谢")
    document.add_paragraph("参考文献")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="paper-toc-order.docx")
    roles = [section.role for section in model.logical_sections]
    assert roles == [
        "abstract",
        "keywords",
        "toc",
        "body",
        "acknowledgements",
        "references",
    ]
    ruleset = RuleSet(
        id="ruleset_structure_order",
        name="结构顺序规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="structure_required_sections",
                category=RuleCategory.structure,
                target=RuleTarget(scope="document.structure", selector="论文结构"),
                expectation={
                    "requiredSections": [
                        "中文摘要",
                        "中文关键词",
                        "目录",
                        "正文",
                        "致谢",
                        "参考文献",
                    ]
                },
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="论文结构要求"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert findings == []


def test_heading_checkers_ignore_numbered_toc_entries(tmp_path: Path) -> None:
    path = tmp_path / "paper-toc-heading.docx"
    document = Document()
    document.add_paragraph("目  录")
    document.styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
    document.add_paragraph("2 系统技术选型 3", style="toc 1")
    document.add_heading("绪论", level=1)
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="paper-toc-heading.docx")
    ruleset = RuleSet(
        id="ruleset_toc_heading",
        name="标题规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="heading1_font",
                category=RuleCategory.heading,
                target=RuleTarget(scope="heading.1", selector="Heading 1"),
                expectation={"fontSizePt": 12},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="一级标题小四"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert all(finding.excerpt != "2 系统技术选型 3" for finding in findings)


def test_structure_finding_includes_document_excerpt(tmp_path: Path) -> None:
    path = tmp_path / "paper-missing.docx"
    document = Document()
    document.add_heading("绪论", level=1)
    document.add_paragraph("正文内容")
    document.save(path)
    model = parse_docx(path, document_id="doc_1", source_filename="paper-missing.docx")
    ruleset = RuleSet(
        id="ruleset_structure_missing",
        name="结构规则",
        source_type=SourceType.manual,
        version="1.0.0",
        created_at="2026-04-26T00:00:00+08:00",
        rules=[
            FormatRule(
                id="structure_required_sections",
                category=RuleCategory.structure,
                target=RuleTarget(scope="document.structure", selector="论文结构"),
                expectation={"requiredSections": ["中文摘要", "中文关键词", "目录", "正文"]},
                severity=Severity.major,
                source=RuleSource(type=SourceType.manual, excerpt="论文结构要求"),
                confidence=1,
            )
        ],
    )

    findings = CheckEngine().run(model, ruleset.rules)

    assert len(findings) == 1
    assert findings[0].excerpt is not None
    assert "绪论" in findings[0].excerpt
