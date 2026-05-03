from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from docchecker.services.requirement_parser import (
    extract_requirement_text,
    parse_requirement_document,
)


def test_extract_requirement_text_includes_word_comments(tmp_path: Path) -> None:
    path = tmp_path / "requirement.docx"
    document = Document()
    document.add_paragraph("中文论文题目")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "正文"
    table.cell(0, 1).text = "宋体小四"
    document.save(path)

    with ZipFile(path, "a", ZIP_DEFLATED) as package:
        package.writestr(
            "word/comments.xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="user">
    <w:p><w:r><w:t>三号宋体加粗居中，一般不多于30个字。</w:t></w:r></w:p>
  </w:comment>
</w:comments>
""",
        )

    text = extract_requirement_text(path)

    assert "中文论文题目" in text
    assert "三号宋体加粗居中" in text
    assert "paragraph:1" in text
    assert "table:1,row:1" in text
    assert "comment:0" in text


def test_parse_requirement_document_keeps_structured_blocks(tmp_path: Path) -> None:
    path = tmp_path / "requirement.docx"
    document = Document()
    document.add_heading("目录", level=1)
    document.add_paragraph("目录自动生成，列至三级标题。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "参考文献"
    table.cell(0, 1).text = "按 GB/T 7714 编排"
    document.save(path)

    model = parse_requirement_document(path)

    assert [block.type for block in model.blocks] == ["paragraph", "paragraph", "table"]
    assert model.blocks[1].heading_path == ["目录"]
    assert model.blocks[2].cells == ["参考文献", "按 GB/T 7714 编排"]
    assert "table:1,row:1" in model.markdown


def test_parse_requirement_document_keeps_exemplar_formatting(tmp_path: Path) -> None:
    path = tmp_path / "requirement-format.docx"
    document = Document()
    paragraph = document.add_paragraph("1 绪论")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)
    document.save(path)

    model = parse_requirement_document(path)
    formatting = model.blocks[0].context["formatting"]

    assert formatting["fontFamilyEastAsia"] == "黑体"
    assert formatting["fontSizePt"] == 16
    assert formatting["alignment"] == "center"


def test_parse_requirement_document_links_comments_to_anchored_paragraph(
    tmp_path: Path,
) -> None:
    path = tmp_path / "requirement-comment-anchor.docx"
    document = Document()
    document.add_paragraph("中文论文题目")
    document.add_paragraph("××××××（正文段落）")
    document.save(path)

    with ZipFile(path, "a", ZIP_DEFLATED) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        document_xml = document_xml.replace(
            "<w:t>××××××（正文段落）</w:t>",
            '<w:commentRangeStart w:id="9"/><w:t>××××××（正文段落）</w:t>'
            '<w:commentRangeEnd w:id="9"/><w:r><w:commentReference w:id="9"/></w:r>',
        )
        package.writestr("word/document.xml", document_xml)
        package.writestr(
            "word/comments.xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="9" w:author="user">
    <w:p><w:r><w:t>中文 宋体、小四，两端对齐，首行缩进2字符。</w:t></w:r></w:p>
  </w:comment>
</w:comments>
""",
        )

    model = parse_requirement_document(path)
    comment = next(block for block in model.blocks if block.location == "comment:9")

    assert comment.context["nearby_location"] == "paragraph:2"
    assert comment.context["nearby_text"] == "××××××（正文段落）"
