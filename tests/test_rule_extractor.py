import json

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from docchecker.core.config import get_settings
from docchecker.domain.enums import SourceType
from docchecker.services.requirement_parser import parse_requirement_document
from docchecker.services.rule_extractor import (
    RuleExtractionConfigurationError,
    extract_rules_from_requirement_document,
    extract_rules_from_text,
)


def test_extract_rules_from_manual_text(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    result = extract_rules_from_text(
        "正文宋体小四，1.5倍行距，首行缩进2字符，页边距上下2.5cm。",
        source_type=SourceType.manual,
    )

    rule_ids = {rule.id for rule in result.rules}

    assert "body_font" in rule_ids
    assert "body_line_spacing" in rule_ids
    assert "body_first_line_indent" in rule_ids
    assert "page_margin_top" in rule_ids
    assert result.extraction_summary.structured_rules >= 4
    assert result.parse_warnings == []


def test_extract_rules_reports_warning_for_unrecognized_text() -> None:
    result = extract_rules_from_text("按学校要求执行。", source_type=SourceType.manual)

    assert result.rules == []
    assert result.parse_warnings
    assert result.extraction_summary.structured_rules == 0


def test_extract_title_rule_from_comment_style_text() -> None:
    result = extract_rules_from_text(
        "三号宋体加粗居中，一般不多于30个字。",
        source_type=SourceType.requirement_doc,
    )

    title_rule = next(rule for rule in result.suggested_rules if rule.id == "title_format")

    assert title_rule.expectation["fontFamilyEastAsia"] == "宋体"
    assert title_rule.expectation["fontSizePt"] == 16
    assert title_rule.expectation["bold"] is True
    assert title_rule.capability_status == "needs_confirmation"
    assert title_rule.confirmation_required is True


def test_needs_confirmation_candidate_does_not_enter_auto_rules(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    result, _ = _extract_with_mocked_llm(
        monkeypatch,
        {
            "rule_candidates": [
                {
                    "category": "font",
                    "target_scope": "body.paragraph",
                    "selector": None,
                    "expectation": {"fontFamilyEastAsia": "宋体"},
                    "evidence_span": "正文一般采用宋体。",
                    "location": "paragraph:1",
                    "checkability": "needs_confirmation",
                    "confidence": 0.72,
                    "reason": "措辞不够确定",
                }
            ]
        },
    )

    assert "font_candidate" not in {rule.id for rule in result.rules}
    suggested = next(rule for rule in result.suggested_rules if rule.id == "font_candidate")
    assert suggested.capability_status == "needs_confirmation"
    assert suggested.confirmation_required is True
    assert result.extraction_summary.needs_confirmation_rules == 1


def test_unsupported_candidate_field_is_reported_not_auto_checked(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    result, _ = _extract_with_mocked_llm(
        monkeypatch,
        {
            "rule_candidates": [
                {
                    "category": "paragraph",
                    "target_scope": "body.paragraph",
                    "selector": None,
                    "expectation": {"spaceBetweenNumberAndTitle": True},
                    "evidence_span": "标题序号和标题之间空1格。",
                    "location": "paragraph:2",
                    "checkability": "checkable",
                    "confidence": 0.9,
                    "reason": None,
                }
            ]
        },
    )

    assert result.rules == []
    assert result.suggested_rules == []
    assert any(
        item.reason_code == "unsupported_field"
        and item.capability_status == "unsupported"
        for item in result.unsupported_requirements
    )


def test_extract_rules_keeps_requirement_locations() -> None:
    result = extract_rules_from_text(
        "paragraph:1\t正文宋体小四。\n"
        "table:1,row:2\t页边距：上2.5cm，下2.0cm，左3cm，右2cm。",
        source_type=SourceType.requirement_doc,
    )

    body_rule = next(rule for rule in result.rules if rule.id == "body_font")
    left_margin = next(rule for rule in result.rules if rule.id == "page_margin_left")

    assert body_rule.source.location == "paragraph:1"
    assert left_margin.source.location == "table:1,row:2"
    assert left_margin.expectation["margin_left_cm"] == 3


def test_extract_rules_maps_semantic_requirements_to_checkable_rules(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    result = extract_rules_from_text(
        "参考文献按 GB/T 7714 编排。\n目录自动生成，列至三级标题。",
        source_type=SourceType.requirement_doc,
    )

    rule_ids = {rule.id for rule in result.rules}

    assert "reference_basic_entries" in rule_ids
    assert "toc_basic_shape" in rule_ids
    assert result.extraction_summary.unsupported_requirements == 0


def test_extract_rules_does_not_treat_caption_or_abstract_notes_as_body_rules(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    result = extract_rules_from_text(
        "\n".join(
            [
                "图注用五号字宋体，两端对齐，首行缩进2字符，段前6磅，段后12磅。",
                "摘要中应避免出现公式、图表，不引用参考文献。",
            ]
        ),
        source_type=SourceType.requirement_doc,
    )

    rule_ids = {rule.id for rule in result.rules}

    assert "body_first_line_indent" not in rule_ids
    assert "body_paragraph_spacing" not in rule_ids
    assert "caption_basic_pattern" in rule_ids
    assert "reference_basic_entries" not in rule_ids


def test_extract_rules_accepts_heading_first_line_indent_with_ge_character(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    result = extract_rules_from_text(
        "一级标题左对齐，首行缩进2个字符，段后6磅。",
        source_type=SourceType.requirement_doc,
    )

    rule = next(rule for rule in result.rules if rule.id == "heading_1_first_line_indent")

    assert rule.target.scope == "heading.1"
    assert rule.expectation == {"firstLineIndentCm": 0.74}


def test_requirement_doc_exemplar_format_overrides_comment_text(
    tmp_path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    path = tmp_path / "requirement-exemplar.docx"
    document = Document()
    paragraph = document.add_paragraph("1 绪论")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16)
    document.add_paragraph(
        "数字 Times New Roman 小四，中文 宋体、小四，"
        "序号和一级标题之间空1格。段落设置中，对齐方式左对齐。"
    )
    document.save(path)

    requirement = parse_requirement_document(path)
    result = extract_rules_from_requirement_document(
        requirement,
        source_type=SourceType.requirement_doc,
    )
    rule = next(rule for rule in result.rules if rule.id == "heading1_font")

    assert rule.expectation["fontFamilyEastAsia"] == "黑体"
    assert rule.expectation["fontSizePt"] == 16
    assert rule.expectation["alignment"] == "center"
    assert rule.confidence == 0.98
    alignment_rule = next(rule for rule in result.rules if rule.id == "heading_1_alignment")
    assert alignment_rule.expectation == {"alignment": "center"}
    assert alignment_rule.confidence == 0.98


def test_body_comment_anchor_extracts_body_font_rule(
    tmp_path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    path = tmp_path / "requirement-body-comment.docx"
    document = Document()
    document.add_paragraph("1 绪论")
    document.add_paragraph("××××××（正文段落）")
    document.save(path)

    from zipfile import ZIP_DEFLATED, ZipFile

    with ZipFile(path, "a", ZIP_DEFLATED) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")
        document_xml = document_xml.replace(
            "<w:t>××××××（正文段落）</w:t>",
            '<w:commentRangeStart w:id="15"/><w:t>××××××（正文段落）</w:t>'
            '<w:commentRangeEnd w:id="15"/><w:r><w:commentReference w:id="15"/></w:r>',
        )
        package.writestr("word/document.xml", document_xml)
        package.writestr(
            "word/comments.xml",
            (
                """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="15" w:author="user">
    <w:p><w:r><w:t>"""
                "中文 宋体、小四，英文和数字 Times New Roman、小四，"
                "两端对齐，首行缩进2字符，行距1.5倍。"
                """</w:t></w:r></w:p>
  </w:comment>
</w:comments>
"""
            ),
        )

    requirement = parse_requirement_document(path)
    result = extract_rules_from_requirement_document(
        requirement,
        source_type=SourceType.requirement_doc,
    )
    body_font = next(rule for rule in result.rules if rule.id == "body_font")
    alignment = next(rule for rule in result.rules if rule.id == "body_alignment")

    assert body_font.target.scope == "body.paragraph"
    assert body_font.target.selector is None
    assert body_font.expectation == {"fontFamilyEastAsia": "宋体", "fontSizePt": 12}
    assert alignment.expectation == {"alignment": "justify"}


def test_body_style_cluster_extracts_repeated_body_exemplar_rule(
    tmp_path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "local")
    get_settings.cache_clear()
    path = tmp_path / "requirement-body-cluster.docx"
    document = Document()
    for text in ["××××××（正文段落）", "××××××××（正文段落）"]:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(24)
        run = paragraph.runs[0]
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)
    document.save(path)

    requirement = parse_requirement_document(path)
    result = extract_rules_from_requirement_document(
        requirement,
        source_type=SourceType.requirement_doc,
    )
    body_font = next(rule for rule in result.rules if rule.id == "body_font")

    assert body_font.source.evidence_type == "style_cluster"
    assert body_font.expectation == {"fontFamilyEastAsia": "宋体", "fontSizePt": 12}


def test_extract_rules_from_golden_corpus() -> None:
    samples = [
        (
            "论文题目三号黑体加粗居中。正文宋体小四，1.5倍行距，首行缩进2字符。",
            {"title_format", "body_font", "body_line_spacing", "body_first_line_indent"},
        ),
        (
            "一级标题黑体三号加粗居中；二级标题黑体小三加粗；三级标题宋体四号加粗。",
            {"heading1_font", "heading2_font", "heading3_font"},
        ),
        (
            "A4纸，页边距上2.5cm，下2.5cm，左3cm，右2cm，页眉1.5cm，页脚1.75cm。",
            {
                "page_size_a4",
                "page_margin_top",
                "page_margin_bottom",
                "page_margin_left",
                "page_margin_right",
                "page_header_distance",
                "page_footer_distance",
            },
        ),
        (
            "摘要宋体小四，关键词宋体小四。正文段前6磅，段后6磅。",
            {"abstract_font", "keywords_font", "body_paragraph_spacing"},
        ),
        (
            "正文宋体小四。图题、表题采用五号宋体；参考文献按学校规范著录。",
            {"body_font"},
        ),
    ]

    for text, expected_rule_ids in samples:
        result = extract_rules_from_text(text, source_type=SourceType.manual)
        rule_ids = {rule.id for rule in result.rules}

        assert expected_rule_ids <= rule_ids


def test_hybrid_mode_requires_explicit_llm_configuration(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "hybrid")
    monkeypatch.setenv("DOC_CHECKER_LLM_API_BASE", "")
    monkeypatch.setenv("DOC_CHECKER_LLM_API_KEY", "")
    monkeypatch.setenv("DOC_CHECKER_LLM_MODEL", "")

    from docchecker.core.config import get_settings

    get_settings.cache_clear()
    try:
        with pytest.raises(RuleExtractionConfigurationError):
            extract_rules_from_text("目录自动生成。", source_type=SourceType.requirement_doc)
    finally:
        get_settings.cache_clear()


def test_hybrid_mode_accepts_schema_compliant_llm_candidates(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    result, request_payload = _extract_with_mocked_llm(
        monkeypatch,
        {
            "rule_candidates": [
                {
                    "category": "structure",
                    "target_scope": "document.structure",
                    "selector": "论文结构",
                    "expectation": {"requiredSections": ["中文摘要", "正文"]},
                    "evidence_span": "论文应包括中文摘要和正文。",
                    "location": "paragraph:42",
                    "checkability": "checkable",
                    "confidence": 0.8,
                    "reason": None,
                }
            ]
        },
    )

    system_prompt = request_payload["messages"][0]["content"]
    assert "category 只能是 page" in system_prompt
    assert "checkability 只能是 checkable" in system_prompt
    assert "expectation 必须是 JSON object" in system_prompt
    assert "structure_required_sections" in {rule.id for rule in result.rules}
    assert result.extraction_trace is not None
    assert result.extraction_trace.issues == []


def test_hybrid_mode_rejects_invalid_llm_candidate_schema(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    result, _ = _extract_with_mocked_llm(
        monkeypatch,
        {
            "rule_candidates": [
                {
                    "category": "abstract_word_count",
                    "target_scope": "abstract",
                    "selector": "中文摘要",
                    "expectation": "中文摘要要求300字左右",
                    "evidence_span": "中文摘要要求300字左右。",
                    "location": "paragraph:42",
                    "checkability": "specific",
                    "confidence": 0.8,
                }
            ]
        },
    )

    assert result.rules == []
    assert result.extraction_trace is not None
    assert len(result.extraction_trace.issues) == 1
    issue = result.extraction_trace.issues[0]
    assert issue.reason_code == "invalid_llm_response"
    assert issue.message == (
        "LLM 返回的规则候选格式不符合系统 schema，"
        "已拒绝本次 LLM 候选；本地规则抽取结果仍可使用。"
    )
    assert issue.excerpt is not None
    assert "abstract_word_count" in issue.excerpt
    assert "validation errors" not in issue.message


def test_hybrid_mode_drops_unsupported_llm_expectation_fields(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    result, _ = _extract_with_mocked_llm(
        monkeypatch,
        {
            "rule_candidates": [
                {
                    "category": "heading",
                    "target_scope": "heading.level1",
                    "selector": "一级标题",
                    "expectation": {
                        "fontSizePt": 12,
                        "spaceBetweenNumberAndTitle": True,
                    },
                    "evidence_span": "一级标题字号12，序号和标题之间空1格。",
                    "location": "paragraph:1",
                    "checkability": "checkable",
                    "confidence": 0.8,
                }
            ]
        },
    )

    rule = next(rule for rule in result.rules if rule.id == "heading_candidate")

    assert rule.expectation == {"fontSizePt": 12}
    assert any(
        issue.reason_code == "unsupported_field"
        and "spaceBetweenNumberAndTitle" in issue.message
        for issue in result.extraction_trace.issues
    )
    assert any(
        item.reason_code == "unsupported_field"
        for item in result.unsupported_requirements
    )


def _extract_with_mocked_llm(
    monkeypatch: pytest.MonkeyPatch,
    llm_content: dict[str, object],
):
    monkeypatch.setenv("DOC_CHECKER_RULE_EXTRACTOR_MODE", "hybrid")
    monkeypatch.setenv("DOC_CHECKER_LLM_API_BASE", "https://llm.example.test")
    monkeypatch.setenv("DOC_CHECKER_LLM_API_KEY", "test-key")
    monkeypatch.setenv("DOC_CHECKER_LLM_MODEL", "test-model")

    from docchecker.core.config import get_settings

    get_settings.cache_clear()
    captured_payload: dict[str, object] = {}

    class FakeResponse:
        def __enter__(self) -> "FakeResponse":
            return self

        def __exit__(self, *args: object) -> None:
            return None

        def read(self) -> bytes:
            return json.dumps(
                {
                    "choices": [
                        {
                            "message": {
                                "content": json.dumps(
                                    llm_content,
                                    ensure_ascii=False,
                                )
                            }
                        }
                    ]
                },
                ensure_ascii=False,
            ).encode("utf-8")

    def fake_urlopen(request: object, timeout: int) -> FakeResponse:
        assert timeout == 60
        captured_payload.update(json.loads(request.data.decode("utf-8")))
        return FakeResponse()

    monkeypatch.setattr(
        "docchecker.services.rule_extractor.urllib.request.urlopen",
        fake_urlopen,
    )
    try:
        result = extract_rules_from_text(
            "中文摘要要求300字左右。",
            source_type=SourceType.requirement_doc,
        )
        return result, captured_payload
    finally:
        get_settings.cache_clear()
