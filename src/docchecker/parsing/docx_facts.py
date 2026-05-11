import re

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from docchecker.domain.document import (
    AbstractFact,
    CaptionFact,
    EffectiveFormatFact,
    FieldFact,
    HeaderFooterFact,
    HeaderFooterParagraphFact,
    LogicalSectionNode,
    NumberingFact,
    ParagraphNode,
    ReferenceEntryFact,
    ReferenceFacts,
    StyleFact,
    TableCellFact,
    TableFact,
    TocFact,
)
from docchecker.parsing.docx_formatting import (
    _alignment_name,
    _effective_format_sources,
    _effective_format_values,
    _font_name,
)
from docchecker.parsing.docx_sections import (
    _is_toc_entry,
    _normalize_section_text,
)


def _header_footer_facts(
    document,
    default_font_name: str | None,
) -> list[HeaderFooterFact]:
    facts: list[HeaderFooterFact] = []
    for section_index, section in enumerate(document.sections):
        for kind, part in [
            ("header.default", section.header),
            ("footer.default", section.footer),
            ("header.first_page", section.first_page_header),
            ("footer.first_page", section.first_page_footer),
            ("header.even_page", section.even_page_header),
            ("footer.even_page", section.even_page_footer),
        ]:
            paragraphs = list(part.paragraphs)
            text = "\n".join(paragraph.text for paragraph in paragraphs if paragraph.text.strip())
            paragraph_facts = [
                HeaderFooterParagraphFact(
                    section_index=section_index,
                    kind=kind,
                    index=paragraph_index,
                    text=paragraph.text,
                    style_name=paragraph.style.name if paragraph.style else None,
                    effective_format=_effective_format_values(paragraph, default_font_name),
                    effective_format_sources=_effective_format_sources(
                        paragraph,
                        default_font_name,
                    ),
                )
                for paragraph_index, paragraph in enumerate(paragraphs)
            ]
            facts.append(
                HeaderFooterFact(
                    section_index=section_index,
                    kind=kind,
                    text=text,
                    inherited=bool(getattr(part, "is_linked_to_previous", False)),
                    paragraph_count=len(paragraphs),
                    paragraphs=paragraph_facts,
                )
            )
    return facts


def _table_facts(document, paragraphs: list[ParagraphNode]) -> list[TableFact]:
    facts: list[TableFact] = []
    anchors = _table_anchor_indices(document)
    for table_index, table in enumerate(document.tables):
        preceding_index, following_index = anchors.get(table_index, (None, None))
        caption_text, caption_position = _table_caption(
            paragraphs,
            preceding_index,
            following_index,
        )
        cells = [
            TableCellFact(
                row_index=row_index,
                column_index=column_index,
                text=cell.text,
            )
            for row_index, row in enumerate(table.rows)
            for column_index, cell in enumerate(row.cells)
        ]
        facts.append(
            TableFact(
                index=table_index,
                row_count=len(table.rows),
                column_count=len(table.columns),
                cells=cells,
                caption_text=caption_text,
                caption_position=caption_position,
                preceding_paragraph_index=preceding_index,
                following_paragraph_index=following_index,
            )
        )
    return facts


def _table_anchor_indices(document) -> dict[int, tuple[int | None, int | None]]:
    anchors: dict[int, tuple[int | None, int | None]] = {}
    table_index = 0
    paragraph_index = 0
    block_order: list[tuple[str, int]] = []
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            block_order.append(("paragraph", paragraph_index))
            paragraph_index += 1
        elif isinstance(item, Table):
            block_order.append(("table", table_index))
            paragraph_index += _table_body_paragraph_count(item)
            table_index += 1
    for block_position, (kind, index) in enumerate(block_order):
        if kind != "table":
            continue
        previous_paragraph = next(
            (
                item_index
                for item_kind, item_index in reversed(block_order[:block_position])
                if item_kind == "paragraph"
            ),
            None,
        )
        next_paragraph = next(
            (
                item_index
                for item_kind, item_index in block_order[block_position + 1 :]
                if item_kind == "paragraph"
            ),
            None,
        )
        anchors[index] = (previous_paragraph, next_paragraph)
    return anchors


def _table_body_paragraph_count(table: Table) -> int:
    return sum(len(cell.paragraphs) for row in table.rows for cell in row.cells)


def _table_caption(
    paragraphs: list[ParagraphNode],
    preceding_index: int | None,
    following_index: int | None,
) -> tuple[str | None, str | None]:
    if preceding_index is not None and _is_table_caption_text(paragraphs[preceding_index].text):
        return paragraphs[preceding_index].text, "before"
    if following_index is not None and _is_table_caption_text(paragraphs[following_index].text):
        return paragraphs[following_index].text, "after"
    return None, None


def _is_table_caption_text(text: str) -> bool:
    return bool(re.match(r"^\s*表\s*\d+(?:[.\-]\d+)*", text))


def _caption_facts(
    paragraphs: list[ParagraphNode],
    tables: list[TableFact],
) -> list[CaptionFact]:
    linked_tables: dict[int, tuple[str, int]] = {}
    for table in tables:
        if table.caption_position == "before" and table.preceding_paragraph_index is not None:
            linked_tables[table.preceding_paragraph_index] = ("before", table.index)
        if table.caption_position == "after" and table.following_paragraph_index is not None:
            linked_tables[table.following_paragraph_index] = ("after", table.index)

    facts: list[CaptionFact] = []
    for paragraph in paragraphs:
        kind = _caption_kind(paragraph.text)
        if kind is None:
            continue
        position, target_index = linked_tables.get(paragraph.index, ("standalone", None))
        facts.append(
            CaptionFact(
                kind=kind,
                text=paragraph.text.strip(),
                paragraph_index=paragraph.index,
                position=position,
                target_index=target_index,
            )
        )
    return facts


def _caption_kind(text: str) -> str | None:
    if re.match(r"^\s*表\s*\d+(?:[.\-]\d+)*", text):
        return "table"
    if re.match(r"^\s*图\s*\d+(?:[.\-]\d+)*", text):
        return "figure"
    return None


def _toc_fact(paragraphs: list[ParagraphNode], fields: list[FieldFact]) -> TocFact:
    title_index = next(
        (
            paragraph.index
            for paragraph in paragraphs
            if _normalize_section_text(paragraph.text) in {"目录", "目次"}
        ),
        None,
    )
    entry_indices = [
        paragraph.index for paragraph in paragraphs if _is_toc_entry(paragraph)
    ]
    return TocFact(
        has_title=title_index is not None,
        has_field=any(field.field_type == "TOC" for field in fields),
        entry_count=len(entry_indices),
        title_paragraph_index=title_index,
        entry_paragraph_indices=entry_indices,
    )


def _numbering_fact(paragraph, index: int) -> NumberingFact | None:
    p_pr = getattr(paragraph._p, "pPr", None)
    num_pr = getattr(p_pr, "numPr", None) if p_pr is not None else None
    if num_pr is None:
        return None
    num_id = getattr(num_pr, "numId", None)
    ilvl = getattr(num_pr, "ilvl", None)
    return NumberingFact(
        paragraph_index=index,
        num_id=str(num_id.val) if num_id is not None else None,
        level=str(ilvl.val) if ilvl is not None else None,
        style_name=paragraph.style.name if paragraph.style else None,
        text=paragraph.text,
    )


def _field_facts(paragraph, index: int) -> list[FieldFact]:
    facts: list[FieldFact] = []
    simple_fields = paragraph._p.findall(".//" + qn("w:fldSimple"))
    for node in simple_fields:
        instruction = (node.get(qn("w:instr")) or "").strip()
        if not instruction:
            continue
        facts.append(
            FieldFact(
                paragraph_index=index,
                field_type=instruction.split()[0].upper(),
                instruction=instruction,
                text=paragraph.text or None,
            )
        )
    for node in paragraph._p.iter(qn("w:instrText")):
        instruction = "".join(node.itertext()).strip()
        if not instruction:
            continue
        field_type = instruction.split()[0].upper()
        facts.append(
            FieldFact(
                paragraph_index=index,
                field_type=field_type,
                instruction=instruction,
                text=paragraph.text or None,
            )
        )
    return facts


def _reference_facts(
    paragraphs: list[ParagraphNode],
    logical_sections: list[LogicalSectionNode],
) -> ReferenceFacts:
    reference_section = next(
        (section for section in logical_sections if section.role == "references"),
        None,
    )
    if reference_section is None:
        candidate_paragraphs = paragraphs
    else:
        end_index = (
            reference_section.end_paragraph_index
            if reference_section.end_paragraph_index is not None
            else paragraphs[-1].index
        )
        candidate_paragraphs = [
            paragraph
            for paragraph in paragraphs
            if reference_section.start_paragraph_index < paragraph.index <= end_index
        ]
    entries = [
        ReferenceEntryFact(
            paragraph_index=paragraph.index,
            number=_reference_number(paragraph.text),
            text=paragraph.text,
        )
        for paragraph in candidate_paragraphs
        if _reference_number(paragraph.text) is not None
    ]
    numbers = [entry.number for entry in entries]
    expected = list(range(1, len(entries) + 1))
    return ReferenceFacts(
        has_section=reference_section is not None,
        section_start_paragraph_index=(
            reference_section.start_paragraph_index if reference_section else None
        ),
        entry_count=len(entries),
        numbering_continuous=bool(entries) and numbers == expected,
        entries=entries,
    )


def _reference_number(text: str) -> int | None:
    match = re.match(r"^\s*\[(\d+)]", text)
    return int(match.group(1)) if match else None


def _abstract_facts(
    paragraphs: list[ParagraphNode],
    logical_sections: list[LogicalSectionNode],
) -> list[AbstractFact]:
    facts: list[AbstractFact] = []
    for section_index, section in enumerate(logical_sections):
        if section.role != "abstract":
            continue
        end_index = (
            section.end_paragraph_index
            if section.end_paragraph_index is not None
            else paragraphs[-1].index
        )
        content_paragraphs = [
            paragraph
            for paragraph in paragraphs
            if section.start_paragraph_index < paragraph.index <= end_index
        ]
        title_paragraph = next(
            (
                paragraph
                for paragraph in paragraphs
                if paragraph.index == section.start_paragraph_index
            ),
            None,
        )
        inline_content = (
            _inline_abstract_content(title_paragraph.text)
            if title_paragraph is not None
            else ""
        )
        keyword = next(
            (
                paragraph.text.strip()
                for paragraph in content_paragraphs
                if _is_keyword_paragraph(paragraph.text)
            ),
            None,
        )
        if keyword is None:
            keyword = _following_keyword_text(paragraphs, logical_sections, section_index)
        content_parts = [inline_content] if inline_content else []
        content_parts.extend(
            paragraph.text.strip()
            for paragraph in content_paragraphs
            if paragraph.text.strip() and not _is_keyword_paragraph(paragraph.text)
        )
        content_text = "\n".join(content_parts)
        facts.append(
            AbstractFact(
                language=_abstract_language(section.title),
                title_paragraph_index=section.start_paragraph_index,
                content_text=content_text,
                word_count=_abstract_word_count(content_text),
                keyword_text=keyword,
                keyword_count=_keyword_count(keyword),
                has_keywords=keyword is not None,
            )
        )
    return facts


def _following_keyword_text(
    paragraphs: list[ParagraphNode],
    logical_sections: list[LogicalSectionNode],
    section_index: int,
) -> str | None:
    if section_index + 1 >= len(logical_sections):
        return None
    next_section = logical_sections[section_index + 1]
    if next_section.role != "keywords":
        return None
    end_index = (
        next_section.end_paragraph_index
        if next_section.end_paragraph_index is not None
        else paragraphs[-1].index
    )
    return next(
        (
            paragraph.text.strip()
            for paragraph in paragraphs
            if next_section.start_paragraph_index <= paragraph.index <= end_index
            and _is_keyword_paragraph(paragraph.text)
        ),
        None,
    )


def _abstract_language(title: str) -> str:
    normalized = title.lower()
    if "英文" in title or "abstract" in normalized:
        return "en"
    return "zh"


def _inline_abstract_content(text: str) -> str:
    stripped = text.strip()
    for prefix in ["摘 要", "摘要", "中文摘要", "Abstract", "英文摘要"]:
        match = re.match(rf"^{re.escape(prefix)}\s*[:：]\s*(.+)$", stripped, flags=re.I)
        if match:
            return match.group(1).strip()
    return ""


def _abstract_word_count(text: str) -> int:
    chinese_chars = re.findall(r"[\u4e00-\u9fff]", text)
    ascii_words = re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?", text)
    return len(chinese_chars) + len(ascii_words)


def _is_keyword_paragraph(text: str) -> bool:
    normalized = _normalize_section_text(text)
    return (
        normalized.startswith("关键词")
        or normalized.startswith("关键字")
        or normalized.startswith("keywords")
        or normalized.startswith("key words")
    )


def _keyword_count(text: str | None) -> int:
    if not text:
        return 0
    values = re.sub(
        r"^\s*(关键词|关键字|keywords|key words)\s*[:：]?",
        "",
        text,
        flags=re.I,
    )
    parts = [part.strip() for part in re.split(r"[;；,，、\s]+", values) if part.strip()]
    return len(parts)


def _effective_format_facts(paragraphs: list[ParagraphNode]) -> list[EffectiveFormatFact]:
    return [
        EffectiveFormatFact(
            owner_type="paragraph",
            owner_id=f"paragraph:{paragraph.index}",
            paragraph_index=paragraph.index,
            section_index=paragraph.section_index,
            style_name=paragraph.style_name,
            values=paragraph.raw.get("effective_format", {}),
            sources=paragraph.raw.get("effective_format_sources", {}),
        )
        for paragraph in paragraphs
    ]


def _style_facts(document) -> list[StyleFact]:
    facts: list[StyleFact] = []
    for style in document.styles:
        if not style.name:
            continue
        paragraph_format = getattr(style, "paragraph_format", None)
        font = getattr(style, "font", None)
        formatting = {
            "fontFamily": _font_name(font) if font is not None else None,
            "fontFamilyEastAsia": _font_name(font, script="east_asia")
            if font is not None
            else None,
            "fontSizePt": font.size.pt if font is not None and font.size is not None else None,
            "bold": font.bold if font is not None else None,
            "alignment": _alignment_name(paragraph_format.alignment)
            if paragraph_format is not None
            else None,
            "spaceBeforePt": paragraph_format.space_before.pt
            if paragraph_format is not None and paragraph_format.space_before is not None
            else None,
            "spaceAfterPt": paragraph_format.space_after.pt
            if paragraph_format is not None and paragraph_format.space_after is not None
            else None,
        }
        facts.append(
            StyleFact(
                name=style.name,
                style_id=style.style_id,
                type=str(style.type),
                base_style=_base_style_name(style),
                formatting={key: value for key, value in formatting.items() if value is not None},
            )
        )
    return facts


def _base_style_name(style) -> str | None:
    base_style = getattr(style, "base_style", None)
    return base_style.name if base_style else None
