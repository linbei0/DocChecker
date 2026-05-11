from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from docchecker.domain.document import (
    DocumentFacts,
    DocumentModel,
    ParagraphNode,
    ParseWarning,
    SectionNode,
)
from docchecker.parsing.docx_facts import (
    _abstract_facts,
    _base_style_name,
    _caption_facts,
    _effective_format_facts,
    _field_facts,
    _header_footer_facts,
    _numbering_fact,
    _reference_facts,
    _style_facts,
    _table_facts,
    _toc_fact,
)
from docchecker.parsing.docx_formatting import (
    ParagraphLocator,
    _alignment_name,
    _bold,
    _document_default_font_name,
    _effective_format_sources,
    _effective_format_values,
    _emu_to_cm,
    _font_family,
    _font_size_pt,
    _line_spacing,
    _paragraph_format_value,
    _run_spans,
    _script_font_families,
    _script_font_family,
    _space_pt,
)
from docchecker.parsing.docx_sections import (
    _build_logical_sections,
    _compact_text,
    _is_heading_style,
    _is_known_section_title,
)


def _iter_body_paragraph_contexts(document) -> list[tuple[Paragraph, ParagraphLocator]]:
    contexts: list[tuple[Paragraph, ParagraphLocator]] = []
    table_index = 0
    for item in document.iter_inner_content():
        if isinstance(item, Paragraph):
            contexts.append((item, _paragraph_locator()))
        elif isinstance(item, Table):
            contexts.extend(_table_paragraph_contexts(item, table_index))
            table_index += 1
    return contexts


def _table_paragraph_contexts(
    table: Table,
    table_index: int,
) -> list[tuple[Paragraph, ParagraphLocator]]:
    contexts: list[tuple[Paragraph, ParagraphLocator]] = []
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            for cell_paragraph_index, paragraph in enumerate(cell.paragraphs):
                contexts.append(
                    (
                        paragraph,
                        _paragraph_locator(
                            table_index=table_index,
                            row_index=row_index,
                            column_index=column_index,
                            cell_paragraph_index=cell_paragraph_index,
                        ),
                    )
                )
            for nested_table in cell.tables:
                contexts.extend(_table_paragraph_contexts(nested_table, table_index))
    return contexts


def _paragraph_locator(
    *,
    table_index: int | None = None,
    row_index: int | None = None,
    column_index: int | None = None,
    cell_paragraph_index: int | None = None,
) -> ParagraphLocator:
    return {
        "story": "body",
        "part_name": "word/document.xml",
        "table_index": table_index,
        "row_index": row_index,
        "column_index": column_index,
        "cell_paragraph_index": cell_paragraph_index,
    }


def parse_docx(path: Path, *, document_id: str, source_filename: str) -> DocumentModel:
    package_parts: list[str]
    with ZipFile(path) as package:
        package_parts = package.namelist()
        image_count = len([name for name in package_parts if name.startswith("word/media/")])
        xml_parts = {
            name: package.read(name).decode("utf-8", errors="replace")
            for name in package_parts
            if name.startswith("word/") and name.endswith(".xml")
        }

    document = Document(path)
    default_font_name = _document_default_font_name(document)
    warnings: list[ParseWarning] = []
    sections = [
        SectionNode(
            index=index,
            page_width_cm=_emu_to_cm(section.page_width),
            page_height_cm=_emu_to_cm(section.page_height),
            margin_top_cm=_emu_to_cm(section.top_margin),
            margin_bottom_cm=_emu_to_cm(section.bottom_margin),
            margin_left_cm=_emu_to_cm(section.left_margin),
            margin_right_cm=_emu_to_cm(section.right_margin),
            header_distance_cm=_emu_to_cm(section.header_distance),
            footer_distance_cm=_emu_to_cm(section.footer_distance),
        )
        for index, section in enumerate(document.sections)
    ]

    body_paragraph_contexts = list(_iter_body_paragraph_contexts(document))
    paragraphs: list[ParagraphNode] = []
    current_section_name: str | None = None
    for index, (paragraph, locator) in enumerate(body_paragraph_contexts):
        style_name = paragraph.style.name if paragraph.style else None
        text = paragraph.text
        if (
            locator["story"] == "body"
            and locator["table_index"] is None
            and (_is_heading_style(style_name) or _is_known_section_title(text))
            and text.strip()
        ):
            current_section_name = _compact_text(text, max_length=40)
        raw = {"section_name": current_section_name} if current_section_name else {}
        raw.update({key: value for key, value in locator.items() if value is not None})
        raw["font_family_east_asia_values"] = _script_font_families(
            paragraph,
            default_font_name,
            script="east_asia",
        )
        raw["font_family_ascii_values"] = _script_font_families(
            paragraph,
            default_font_name,
            script="ascii",
        )
        raw["effective_format"] = _effective_format_values(paragraph, default_font_name)
        raw["effective_format_sources"] = _effective_format_sources(
            paragraph,
            default_font_name,
        )
        runs = _run_spans(paragraph, default_font_name)
        paragraphs.append(
            ParagraphNode(
                index=index,
                text=text,
                story=str(locator["story"]),
                part_name=str(locator["part_name"]),
                style_name=style_name,
                table_index=locator["table_index"],
                row_index=locator["row_index"],
                column_index=locator["column_index"],
                cell_paragraph_index=locator["cell_paragraph_index"],
                font_family=_font_family(paragraph, default_font_name),
                font_family_east_asia=_script_font_family(
                    paragraph,
                    default_font_name,
                    script="east_asia",
                ),
                font_family_ascii=_script_font_family(
                    paragraph,
                    default_font_name,
                    script="ascii",
                ),
                font_size_pt=_font_size_pt(paragraph),
                bold=_bold(paragraph),
                alignment=_alignment_name(_paragraph_format_value(paragraph, "alignment")),
                first_line_indent_cm=_emu_to_cm(
                    _paragraph_format_value(paragraph, "first_line_indent")
                ),
                line_spacing=_line_spacing(paragraph),
                space_before_pt=_space_pt(paragraph, "space_before"),
                space_after_pt=_space_pt(paragraph, "space_after"),
                runs=runs,
                raw=raw,
            )
        )

    styles = {
        style.name: {
            "style_id": style.style_id,
            "type": str(style.type),
            "base_style": _base_style_name(style),
        }
        for style in document.styles
        if style.name
    }
    if not paragraphs:
        warnings.append(ParseWarning(code="empty_document", message="文档没有可解析段落。"))

    logical_sections = _build_logical_sections(paragraphs)
    table_facts = _table_facts(document, paragraphs)
    field_facts = [
        fact
        for index, (paragraph, _locator) in enumerate(body_paragraph_contexts)
        for fact in _field_facts(paragraph, index)
    ]

    facts = DocumentFacts(
        xml_parts=xml_parts,
        headers_footers=_header_footer_facts(document, default_font_name),
        tables=table_facts,
        captions=_caption_facts(paragraphs, table_facts),
        toc=_toc_fact(paragraphs, field_facts),
        numbering=[
            fact
            for index, (paragraph, _locator) in enumerate(body_paragraph_contexts)
            if (fact := _numbering_fact(paragraph, index)) is not None
        ],
        fields=field_facts,
        styles=_style_facts(document),
        effective_formats=_effective_format_facts(paragraphs),
        references=_reference_facts(paragraphs, logical_sections),
        abstracts=_abstract_facts(paragraphs, logical_sections),
    )

    return DocumentModel(
        document_id=document_id,
        source_filename=source_filename,
        package_parts=package_parts,
        sections=sections,
        logical_sections=logical_sections,
        paragraphs=paragraphs,
        table_count=len(document.tables),
        image_count=image_count,
        styles=styles,
        facts=facts,
        parse_warnings=warnings,
    )
