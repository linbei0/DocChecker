from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree

from docchecker.domain.requirements import RequirementBlock, RequirementDocumentModel

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def parse_requirement_document(path: Path) -> RequirementDocumentModel:
    document = Document(path)
    blocks: list[RequirementBlock] = []
    warnings: list[str] = []
    heading_path: list[str] = []
    paragraph_index = 0
    table_index = 0

    for item in _iter_document_content(document):
        if isinstance(item, Paragraph):
            paragraph_index += 1
            text = _normalize_text(item.text)
            if not text:
                continue
            style_name = item.style.name if item.style else None
            heading_path = _updated_heading_path(heading_path, style_name, text)
            blocks.append(
                RequirementBlock(
                    id=f"paragraph:{paragraph_index}",
                    type="paragraph",
                    location=f"paragraph:{paragraph_index}",
                    text=text,
                    style_name=style_name,
                    heading_path=heading_path.copy(),
                )
            )
        elif isinstance(item, Table):
            table_index += 1
            blocks.extend(_table_blocks(item, table_index, heading_path))

    blocks.extend(_header_footer_blocks(document, heading_path))
    blocks.extend(_extract_comment_blocks(path, blocks))
    if not blocks:
        warnings.append("规范文档没有可解析文本块。")

    return RequirementDocumentModel(
        source_filename=path.name,
        blocks=blocks,
        markdown=_blocks_to_markdown(blocks),
        parse_warnings=warnings,
    )


def extract_requirement_text(path: Path) -> str:
    model = parse_requirement_document(path)
    return "\n".join(f"{block.location}\t{block.text}" for block in model.blocks if block.text)


def _iter_document_content(document):
    iterator = getattr(document, "iter_inner_content", None)
    if iterator:
        yield from iterator()
        return
    yield from document.paragraphs
    yield from document.tables


def _updated_heading_path(
    current: list[str],
    style_name: str | None,
    text: str,
) -> list[str]:
    level = _heading_level(style_name, text)
    if level is None:
        return current
    next_path = current[: level - 1]
    next_path.append(text[:80])
    return next_path


def _heading_level(style_name: str | None, text: str) -> int | None:
    normalized_style = (style_name or "").strip().lower()
    if normalized_style.startswith("heading"):
        suffix = normalized_style.replace("heading", "").strip()
        return int(suffix) if suffix.isdigit() else 1
    if normalized_style.startswith("标题"):
        suffix = normalized_style.replace("标题", "").strip()
        return int(suffix) if suffix.isdigit() else 1
    if text in {"目录", "目 录", "摘要", "Abstract", "参考文献", "致谢"}:
        return 1
    if _starts_with_numbered_heading(text):
        return min(text.split()[0].count(".") + 1, 4)
    return None


def _starts_with_numbered_heading(text: str) -> bool:
    first = text.split(maxsplit=1)[0] if text.split() else ""
    return bool(first and first[0].isdigit() and any(char.isdigit() for char in first))


def _table_blocks(
    table: Table,
    table_index: int,
    heading_path: list[str],
) -> list[RequirementBlock]:
    blocks: list[RequirementBlock] = []
    for row_index, row in enumerate(table.rows, start=1):
        cells = [_normalize_text(cell.text) for cell in row.cells]
        meaningful_cells = [cell for cell in cells if cell]
        if not meaningful_cells:
            continue
        blocks.append(
            RequirementBlock(
                id=f"table:{table_index},row:{row_index}",
                type="table",
                location=f"table:{table_index},row:{row_index}",
                text=" | ".join(meaningful_cells),
                heading_path=heading_path.copy(),
                table_index=table_index,
                row_index=row_index,
                column_count=len(cells),
                cells=meaningful_cells,
            )
        )
    return blocks


def _header_footer_blocks(document, heading_path: list[str]) -> list[RequirementBlock]:
    blocks: list[RequirementBlock] = []
    seen: set[tuple[str, str]] = set()
    for section_index, section in enumerate(document.sections, start=1):
        for block_type, container in [
            ("header", section.header),
            ("footer", section.footer),
        ]:
            for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
                text = _normalize_text(paragraph.text)
                if not text:
                    continue
                key = (block_type, text)
                if key in seen:
                    continue
                seen.add(key)
                blocks.append(
                    RequirementBlock(
                        id=f"{block_type}:{section_index},paragraph:{paragraph_index}",
                        type=block_type,
                        location=f"{block_type}:{section_index},paragraph:{paragraph_index}",
                        text=text,
                        heading_path=heading_path.copy(),
                    )
                )
    return blocks


def _extract_comment_blocks(
    path: Path,
    context_blocks: list[RequirementBlock],
) -> list[RequirementBlock]:
    with ZipFile(path) as package:
        if "word/comments.xml" not in package.namelist():
            return []
        comments_xml = package.read("word/comments.xml")

    root = etree.fromstring(comments_xml)
    comments: list[RequirementBlock] = []
    for comment_index, comment in enumerate(
        root.xpath("//w:comment", namespaces=WORD_NS),
        start=1,
    ):
        text = _normalize_text("".join(comment.xpath(".//w:t/text()", namespaces=WORD_NS)))
        if not text:
            continue
        comment_id = comment.get(f"{{{WORD_NS['w']}}}id") or str(comment_index)
        nearby = _nearby_comment_context(context_blocks, text)
        comments.append(
            RequirementBlock(
                id=f"comment:{comment_id}",
                type="comment",
                location=f"comment:{comment_id}",
                text=text,
                heading_path=nearby.heading_path if nearby else [],
                context={"nearby_location": nearby.location, "nearby_text": nearby.text}
                if nearby
                else {},
            )
        )
    return comments


def _nearby_comment_context(
    blocks: list[RequirementBlock],
    comment_text: str,
) -> RequirementBlock | None:
    for block in blocks:
        if block.text and (block.text in comment_text or comment_text in block.text):
            return block
    return blocks[-1] if blocks else None


def _blocks_to_markdown(blocks: list[RequirementBlock]) -> str:
    lines: list[str] = []
    for block in blocks:
        prefix = " > ".join(block.heading_path)
        title = f"{block.location} [{block.type}]"
        if prefix:
            title = f"{title} {prefix}"
        lines.extend([f"- {title}", f"  {block.text}"])
    return "\n".join(lines)


def _normalize_text(text: str) -> str:
    return " ".join(text.split())
